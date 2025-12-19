"""
Shared Admin API Endpoints for Chatbots

This module exposes consistent REST endpoints (files, indexing, API keys, logs)
that each chatbot can register to integrate with the Admin Portal. It also
includes helpers for loading BYOK secrets from disk into environment vars.
"""

from __future__ import annotations

import json
import os
import uuid
from datetime import datetime, timedelta
from typing import Dict, List, Optional

from flask import jsonify, request

import io
import os
import tempfile
from docx import Document as DocxDocument
from fpdf import FPDF

API_PROVIDER_ENV_MAP = {
    "openai": "OPENAI_API_KEY",
    "azure_openai": "AZURE_OPENAI_API_KEY",
    "anthropic": "ANTHROPIC_API_KEY",
    "google": "GOOGLE_API_KEY",
}

__all__ = [
    "register_admin_endpoints",
    "apply_client_api_keys",
    "load_api_key_for_provider",
    "get_active_provider_for_client",
    "save_conversation_to_json",
    "save_to_report_db",
]


def _log(logger, level: str, message: str):
    if logger:
        getattr(logger, level)(message)
    else:
        print(f"[{level.upper()}] {message}")


def _client_root(config: Dict, client_id: str) -> str:
    root_dir = config.get("ROOT_DIR", "Data")
    # Convert relative paths to absolute paths based on the script location
    if not os.path.isabs(root_dir):
        script_dir = os.path.dirname(os.path.abspath(__file__))
        root_dir = os.path.join(script_dir, root_dir)
    return root_dir


def _ensure_secrets_path(root_dir: str, client_id: str) -> str:
    secrets_dir = os.path.join(root_dir, client_id, "secrets")
    os.makedirs(secrets_dir, exist_ok=True)
    return os.path.join(secrets_dir, "api_keys.json")


def _load_key_store(path: str) -> Dict[str, Dict]:
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as handle:
                data = json.load(handle)
                if isinstance(data, dict):
                    return data
        except Exception:
            pass
    return {}


def _save_key_store(path: str, data: Dict[str, Dict]) -> None:
    tmp_path = f"{path}.tmp"
    with open(tmp_path, "w", encoding="utf-8") as handle:
        json.dump(data, handle, indent=2)
    os.replace(tmp_path, path)


def _save_indexing_history(client_id: str, indexing_mode: str, urls: str = None, sitemap: str = None, status: str = "unknown", client_configs=None) -> None:
    try:
        # Get the client's data directory
        config = client_configs.get(client_id, {}) if client_configs else {}
        root_dir = _client_root(config, client_id)
        client_dir = os.path.join(root_dir, client_id)
        os.makedirs(client_dir, exist_ok=True)
        history_file = os.path.join(client_dir, "indexing_history.json")

        # Load existing history
        if os.path.exists(history_file):
            with open(history_file, 'r', encoding='utf-8') as f:
                history = json.load(f)
        else:
            history = []

        # Add new entry
        history_entry = {
            "timestamp": datetime.now().isoformat(),
            "indexing_mode": indexing_mode,
            "status": status
        }

        if urls:
            history_entry["urls"] = urls
        if sitemap:
            history_entry["sitemap"] = sitemap

        history.append(history_entry)

        # Save back to file (keep last 100 entries)
        with open(history_file, 'w', encoding='utf-8') as f:
            json.dump(history[-100:], f, indent=2)

    except Exception as exc:
        print(f"[ERROR] Error saving indexing history for {client_id}: {exc}")


def _mask_key(key_value: str) -> str:
    if not key_value:
        return ""
    if len(key_value) <= 8:
        return "*" * len(key_value)
    return f"{key_value[:4]}{'*' * (len(key_value) - 6)}{key_value[-2:]}"


def _apply_env(provider_key: str, entry: Dict, logger=None):
    value = entry.get("key_value")
    if not value:
        return False
    env_var = entry.get("env_var") or API_PROVIDER_ENV_MAP.get(provider_key.lower())
    if env_var:
        os.environ[env_var] = value
        _log(logger, "info", f"Loaded API key for provider '{provider_key}' into {env_var}")
        return True
    return False


def apply_client_api_keys(client_id: str, client_configs: Dict, logger=None) -> List[str]:
    """
    Load and apply (set env vars) for all providers saved for a client.
    """
    config = client_configs.get(client_id)
    if not config:
        return []
    path = _ensure_secrets_path(_client_root(config, client_id), client_id)
    store = _load_key_store(path)
    applied = []
    for provider_key, entry in store.items():
        # Safety check: ensure entry is a dict (handle corrupted legacy format)
        if not isinstance(entry, dict):
            continue

        if _apply_env(provider_key, entry, logger):
            applied.append(entry.get("provider") or provider_key)
    return applied


def get_active_provider_for_client(client_id: str, client_configs: Dict) -> Optional[str]:
    """
    Get the provider name of the active API key for a client.
    Returns the first valid provider found in the key store.
    """
    config = client_configs.get(client_id)
    if not config:
        return None
    path = _ensure_secrets_path(_client_root(config, client_id), client_id)
    store = _load_key_store(path)
    
    for provider_key, entry in store.items():
        if not provider_key.startswith("_") and entry.get("key_value"):
            return entry.get("provider") or provider_key
    return None


def load_api_key_for_provider(
    root_dir: str,
    client_id: str,
    provider: str,
    env_var: Optional[str] = None,
    logger=None,
) -> Optional[str]:
    """
    Helper for setup scripts to load a specific provider key from disk.
    """
    provider_key = provider.lower()
    path = _ensure_secrets_path(root_dir, client_id)
    entry = _load_key_store(path).get(provider_key)
    if not entry:
        return None
    key_value = entry.get("key_value")
    if not key_value:
        return None
    env_name = env_var or entry.get("env_var") or API_PROVIDER_ENV_MAP.get(provider_key)
    if env_name:
        os.environ[env_name] = key_value
        _log(logger, "info", f"Loaded {provider} API key into {env_name}")
    return key_value


def save_conversation_to_json(client_id: str, session_id: str, conversation_dict: Dict, client_configs: Dict, logger=None):
    """
    Save a conversation to the chatbot's SQLite database using its existing Log_sql module.

    Args:
        client_id: The client identifier (e.g., "lollypop_design")
        session_id: Unique session identifier
        conversation_dict: The conversation data (questions and responses)
        client_configs: Client configuration dictionary
        logger: Optional logger instance
    """
    try:
        # Import the chatbot's database module dynamically
        # This assumes Log_sql.py is in utils/ folder in the chatbot
        import sys
        import importlib.util

        config = client_configs.get(client_id, {})
        root_dir = config.get("ROOT_DIR", "Data")

        # Try to import the chatbot's Log_sql module
        log_sql_path = os.path.join(os.path.dirname(__file__), "utils", "Log_sql.py")

        if os.path.exists(log_sql_path):
            spec = importlib.util.spec_from_file_location("Log_sql", log_sql_path)
            Log_sql = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(Log_sql)

            # Use the chatbot's database function to save/update conversation
            Log_sql.update_activity_for_session(client_id, session_id, conversation_dict)

            if logger:
                logger.debug(f"Saved conversation to database for session {session_id}")
        else:
            # Fallback: Save directly to SQLite if Log_sql.py not found
            _save_conversation_to_db_direct(client_id, session_id, conversation_dict, config, logger)

    except Exception as e:
        if logger:
            logger.error(f"Error saving conversation to database: {e}")
        else:
            print(f"Error saving conversation to database: {e}")


def _save_conversation_to_db_direct(client_id: str, session_id: str, conversation_dict: Dict, config: Dict, logger=None):
    """
    Direct database save as fallback if Log_sql module not available.

    This function directly saves to the chatbot's Log.db database.
    """
    import sqlite3

    try:
        # Get database path from config or use default
        log_db_path = config.get("LOG_DB_PATH")
        if not log_db_path:
            # Try relative to current directory first (most common)
            log_db_path = os.path.join("application_db", "log_db")
            if not os.path.exists(log_db_path):
                # Try under Data directory
                root_dir = config.get("ROOT_DIR", "Data")
                log_db_path = os.path.join(root_dir, "application_db", "log_db")

        os.makedirs(log_db_path, exist_ok=True)
        log_db_file = os.path.join(log_db_path, "Log.db")

        conn = sqlite3.connect(log_db_file)
        cursor = conn.cursor()

        # Create table if it doesn't exist
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS client_sessions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            client_id TEXT NOT NULL,
            session_id TEXT NOT NULL,
            conversation JSON,
            time TEXT,
            date TEXT,
            summary BOOLEAN DEFAULT 0
        )
        """)

        # Check if session exists
        cursor.execute("SELECT conversation FROM client_sessions WHERE session_id = ?", (session_id,))
        result = cursor.fetchone()

        timestamp = datetime.now()
        time_str = timestamp.strftime("%H:%M:%S")
        date_str = timestamp.strftime("%Y-%m-%d")

        if result:
            # Update existing session
            existing_json = json.loads(result[0])
            for key, value in conversation_dict.items():
                if key in existing_json and isinstance(existing_json[key], dict) and isinstance(value, dict):
                    existing_json[key].update(value)
                else:
                    existing_json[key] = value

            cursor.execute("""
            UPDATE client_sessions
            SET conversation = ?, time = ?, date = ?
            WHERE session_id = ?
            """, (json.dumps(existing_json), time_str, date_str, session_id))
        else:
            # Insert new session
            cursor.execute("""
            INSERT INTO client_sessions (client_id, session_id, conversation, time, date, summary)
            VALUES (?, ?, ?, ?, ?, ?)
            """, (client_id, session_id, json.dumps(conversation_dict), time_str, date_str, 0))

        conn.commit()
        conn.close()

        if logger:
            logger.debug(f"Saved conversation to database (direct) for session {session_id}")

    except Exception as e:
        if logger:
            logger.error(f"Error in direct database save: {e}")
        else:
            print(f"Error in direct database save: {e}")


def save_to_report_db(client_id: str, session_id: str, conversation_dict: Dict, client_configs: Dict, logger=None):
    """
    Stub function for backwards compatibility.

    This function is called by older chatbot code but is no longer used.
    Reports are now generated on-demand via the Admin Portal API using OpenAI.

    The chatbot just needs to save conversations to JSON files, and the
    Admin Portal will extract information when reports are requested.
    """
    # Do nothing - reports are generated on-demand in Admin Portal
    pass


def register_admin_endpoints(app, client_configs, logger=None):
    """
    Register admin API endpoints for PDFs, indexing, and BYOK secrets.
    """

    def log_info(message: str):
        _log(logger, "info", message)

    def log_error(message: str):
        _log(logger, "exception", message)

    def _convert_office_to_pdf_bytes(file_bytes: bytes, ext: str) -> Optional[bytes]:
        """
        Convert DOC/DOCX bytes to a simple PDF (text-only) for upload.
        Returns PDF bytes on success, or None on failure.
        """
        ext = ext.lower()
        text_content = ""

        if ext == ".docx":
            doc = DocxDocument(io.BytesIO(file_bytes))
            text_content = "\n\n".join(p.text for p in doc.paragraphs)
        elif ext == ".doc":
            try:
                import textract

                with tempfile.NamedTemporaryFile(suffix=".doc", delete=True) as tmp:
                    tmp.write(file_bytes)
                    tmp.flush()
                    extracted = textract.process(tmp.name)
                    text_content = extracted.decode("utf-8", errors="ignore")
            except Exception:
                return None
        else:
            return None

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        page_width = pdf.w - 2 * pdf.l_margin # Calculate usable page width
        
        # Iterate over paragraphs (assuming paragraphs are separated by double newlines or similar)
        # Or, just feed the entire text content as one block and let multi_cell handle wrapping
        safe_text_content = text_content.encode("latin-1", "replace").decode("latin-1")
        pdf.multi_cell(page_width, 6, safe_text_content) # Use page_width and reduced line height

        return pdf.output()

    def _validate_client(client_id: str):
        if not client_id or client_id not in client_configs:
            raise ValueError(f"Invalid client_id: {client_id}")
        return client_configs[client_id]

    def _extract_text_from_file(file_bytes: bytes, ext: str) -> str:
        """
        Extract text content from a file for content comparison.
        For DOC/DOCX, converts to PDF first then extracts (matches stored format).
        Returns extracted text or empty string on failure.
        """
        import fitz  # PyMuPDF
        ext = ext.lower()
        text_content = ""
        
        try:
            if ext == ".pdf":
                # Direct PDF extraction
                pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
                for page in pdf_doc:
                    text_content += page.get_text()
                pdf_doc.close()
            elif ext in [".docx", ".doc"]:
                # Convert DOC/DOCX to PDF first, then extract from PDF
                # This ensures we compare "apples to apples" with stored PDFs
                pdf_bytes = _convert_office_to_pdf_bytes(file_bytes, ext)
                if pdf_bytes:
                    pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                    for page in pdf_doc:
                        text_content += page.get_text()
                    pdf_doc.close()
                else:
                    # Fallback to direct extraction if conversion fails
                    if ext == ".docx":
                        doc = DocxDocument(io.BytesIO(file_bytes))
                        text_content = "\n".join(p.text for p in doc.paragraphs)
                    elif ext == ".doc":
                        import textract
                        with tempfile.NamedTemporaryFile(suffix=".doc", delete=True) as tmp:
                            tmp.write(file_bytes)
                            tmp.flush()
                            extracted = textract.process(tmp.name)
                            text_content = extracted.decode("utf-8", errors="ignore")
        except Exception as e:
            _log(logger, "warning", f"Text extraction failed for {ext}: {e}")
            return ""
        
        return text_content.strip()

    def _get_content_hash(text: str) -> str:
        """
        Generate MD5 hash of normalized text content for comparison.
        Normalizes text to handle DOC→PDF conversion differences.
        """
        import hashlib
        import re
        
        # Normalize text for consistent comparison
        normalized = text.lower()                    # Lowercase
        normalized = re.sub(r'\s+', ' ', normalized) # Collapse whitespace
        normalized = normalized.strip()              # Trim ends
        
        return hashlib.md5(normalized.encode('utf-8')).hexdigest()


    @app.route("/api/files/upload", methods=["POST"])
    def upload_file():
        filepath = None  # Track filepath for cleanup on failure
        try:
            client_id = request.form.get("client_id")
            config = _validate_client(client_id)

            if "file" not in request.files:
                return jsonify({"error": "No file provided"}), 400

            upload = request.files["file"]
            if not upload.filename:
                return jsonify({"error": "No file name provided"}), 400

            filename_lower = upload.filename.lower()
            _, ext = os.path.splitext(filename_lower)

            # Accept PDF, DOC, DOCX
            if ext not in [".pdf", ".doc", ".docx"]:
                return jsonify({"error": "Only PDF, DOC, or DOCX files are allowed"}), 400

            root_dir = _client_root(config, client_id)
            upload_dir = os.path.join(root_dir, client_id, "uploads")
            os.makedirs(upload_dir, exist_ok=True)

            file_id = str(uuid.uuid4())

            if ext in [".doc", ".docx"]:
                # Convert to PDF before saving
                file_bytes = upload.read()
                pdf_bytes = _convert_office_to_pdf_bytes(file_bytes, ext)
                if pdf_bytes is None:
                    return jsonify({"error": "DOC/DOCX conversion failed; please upload a PDF"}), 400

                pdf_name = f"{file_id}_{os.path.splitext(upload.filename)[0]}.pdf"
                filepath = os.path.join(upload_dir, pdf_name)
                with open(filepath, "wb") as f:
                    f.write(pdf_bytes)

                saved_filename = pdf_name
            else:
                # Already PDF – save as-is
                saved_filename = f"{file_id}_{upload.filename}"
                filepath = os.path.join(upload_dir, saved_filename)
                upload.save(filepath)

            log_info(f"File uploaded for {client_id}: {saved_filename}")
            return (
                jsonify(
                    {
                        "message": "File uploaded successfully",
                        "document_id": file_id,
                        "filename": upload.filename,
                        "filepath": filepath,
                    }
                ),
                201,
            )
        except ValueError as exc:
            # Cleanup partial file on validation error
            if filepath and os.path.exists(filepath):
                try:
                    os.remove(filepath)
                except Exception:
                    pass
            return jsonify({"error": str(exc)}), 400
        except Exception as exc:
            # Cleanup partial file on any error
            if filepath and os.path.exists(filepath):
                try:
                    os.remove(filepath)
                    log_info(f"Cleaned up partial file: {filepath}")
                except Exception as cleanup_err:
                    log_error(f"Failed to cleanup partial file {filepath}: {cleanup_err}")
            log_error(f"Error uploading file: {exc}")
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/files", methods=["GET"])
    def list_files():
        try:
            client_id = request.args.get("client_id")
            config = _validate_client(client_id)
            root_dir = _client_root(config, client_id)
            upload_dir = os.path.join(root_dir, client_id, "uploads")

            if not os.path.exists(upload_dir):
                return jsonify({"documents": []}), 200

            documents = []
            for filename in os.listdir(upload_dir):
                if not filename.lower().endswith(".pdf"):
                    continue
                filepath = os.path.join(upload_dir, filename)
                file_id = filename.split("_", 1)[0]
                original = filename.split("_", 1)[1] if "_" in filename else filename
                documents.append(
                    {
                        "document_id": file_id,
                        "filename": original,
                        "uploaded_at": os.path.getctime(filepath),
                    }
                )

            return jsonify({"documents": documents}), 200
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except Exception as exc:
            log_error(f"Error listing files: {exc}")
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/files/<document_id>", methods=["DELETE"])
    def delete_file(document_id):
        import gc
        import time
        
        try:
            client_id = request.args.get("client_id")
            config = _validate_client(client_id)
            root_dir = _client_root(config, client_id)
            upload_dir = os.path.join(root_dir, client_id, "uploads")

            if not os.path.exists(upload_dir):
                return jsonify({"error": "No uploads found"}), 404

            for filename in os.listdir(upload_dir):
                if filename.startswith(f"{document_id}_"):
                    filepath = os.path.join(upload_dir, filename)
                    
                    # Force garbage collection to release any file handles
                    gc.collect()
                    
                    # Retry logic for Windows file locking
                    max_retries = 5
                    for attempt in range(max_retries):
                        try:
                            os.remove(filepath)
                            log_info(f"Deleted file for {client_id}: {filename}")
                            return jsonify({"message": "File deleted successfully"}), 200
                        except PermissionError as e:
                            if attempt < max_retries - 1:
                                time.sleep(1.0)  # Wait longer and retry
                                gc.collect()
                            else:
                                log_error(f"Permission denied after {max_retries} attempts: {filepath}")
                                return jsonify({"error": f"File is locked, please try again: {str(e)}"}), 500

            return jsonify({"error": "File not found"}), 404
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except Exception as exc:
            log_error(f"Error deleting file: {exc}")
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/files/check-duplicate", methods=["POST"])
    def check_duplicate_file():
        """
        Check if a file with the same name already exists and compare content.
        Returns duplicate info including whether content is same or different.
        
        Accepts either:
        - JSON: {"client_id": "...", "filename": "..."} (filename check only)
        - Form data: client_id, file (with file content for comparison)
        """
        try:
            # Handle both JSON and form data
            if request.content_type and 'multipart/form-data' in request.content_type:
                client_id = request.form.get("client_id")
                upload = request.files.get("file")
                if upload:
                    filename = upload.filename
                    uploaded_bytes = upload.read()
                    upload.seek(0)  # Reset for potential later use
                else:
                    filename = request.form.get("filename", "").strip()
                    uploaded_bytes = None
            else:
                data = request.json or {}
                client_id = data.get("client_id")
                filename = data.get("filename", "").strip()
                uploaded_bytes = None
            
            config = _validate_client(client_id)
            root_dir = _client_root(config, client_id)
            upload_dir = os.path.join(root_dir, client_id, "uploads")
            
            if not filename:
                return jsonify({"error": "Filename is required"}), 400
            
            # Get uploaded file extension
            _, uploaded_ext = os.path.splitext(filename.lower())
            
            # Normalize filename: DOC/DOCX are stored as PDF
            base_filename = filename.lower()
            name_without_ext, ext = os.path.splitext(base_filename)
            if ext in [".doc", ".docx"]:
                base_filename = name_without_ext + ".pdf"
            
            # Extract text and hash from uploaded file (if provided)
            uploaded_hash = None
            if uploaded_bytes:
                uploaded_text = _extract_text_from_file(uploaded_bytes, uploaded_ext)
                if uploaded_text:
                    uploaded_hash = _get_content_hash(uploaded_text)
            
            duplicates = []
            if os.path.exists(upload_dir):
                for existing_file in os.listdir(upload_dir):
                    # Files are stored as UUID_originalname.pdf
                    parts = existing_file.split("_", 1)
                    if len(parts) > 1:
                        original_name = parts[1].lower()
                        file_id = parts[0]
                    else:
                        original_name = existing_file.lower()
                        file_id = None
                    
                    if original_name == base_filename:
                        filepath = os.path.join(upload_dir, existing_file)
                        
                        # Compare content if we have uploaded file bytes
                        is_same_content = None
                        existing_hash = None
                        if uploaded_hash:
                            try:
                                with open(filepath, "rb") as f:
                                    existing_bytes = f.read()
                                existing_text = _extract_text_from_file(existing_bytes, ".pdf")
                                if existing_text:
                                    existing_hash = _get_content_hash(existing_text)
                                    is_same_content = (uploaded_hash == existing_hash)
                            except Exception as e:
                                log_error(f"Content comparison failed: {e}")
                        
                        duplicates.append({
                            "document_id": file_id,
                            "filename": existing_file,
                            "original_name": parts[1] if len(parts) > 1 else existing_file,
                            "uploaded_at": os.path.getctime(filepath),
                            "size": os.path.getsize(filepath),
                            "is_same_content": is_same_content,
                            "existing_hash": existing_hash
                        })
            
            log_info(f"Duplicate check for {filename}: found {len(duplicates)} duplicate(s)")
            return jsonify({
                "is_duplicate": len(duplicates) > 0,
                "duplicates": duplicates,
                "filename_checked": filename,
                "normalized_filename": base_filename,
                "uploaded_hash": uploaded_hash
            }), 200
            
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except Exception as exc:
            log_error(f"Error checking for duplicates: {exc}")
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/indexing/start", methods=["POST"])
    def start_indexing():
        """
        Start indexing job for a client (PDF / website / sitemap / URLs).
        Mirrors the behavior documented in Admin Portal.
        """
        try:
            data = request.json or {}
            client_id = data.get("client_id")
            index_type = data.get("index_type", "all")
            urls = data.get("urls")
            sitemap = data.get("sitemap")

            log_info(f"[DEBUG] Received indexing request: client_id={client_id}, index_type={index_type}, urls={urls}, sitemap={sitemap}")
            _validate_client(client_id)

            import subprocess
            import sys
            import threading

            job_id = str(uuid.uuid4())
            app.indexing_jobs = getattr(app, "indexing_jobs", {})

            if index_type == "sitemap" and sitemap:
                indexing_mode = f"Sitemap ({sitemap})"
                indexing_mode_detail = f"Sitemap mode: {sitemap}"
            elif index_type == "website" and urls:
                indexing_mode = "Specific URLs"
                indexing_mode_detail = f"URL mode: {urls}"
            elif index_type == "website":
                indexing_mode = "Full Website"
                indexing_mode_detail = "Full website mode"
            elif index_type == "pdf":
                indexing_mode = "PDF"
                indexing_mode_detail = "PDF mode"
            else:
                indexing_mode = "Full"
                indexing_mode_detail = "Full (default) mode"

            app.indexing_jobs[job_id] = {
                "status": "processing",
                "progress": 0,
                "message": f"Starting {indexing_mode} indexing...",
                "indexing_mode": indexing_mode_detail,
                "client_id": client_id,
            }
            log_info(f"[Index] Created job {job_id} with mode: {indexing_mode_detail}")

            def run_indexing():
                try:
                    chatbot_root = os.getcwd()
                    setup_path = os.path.join(chatbot_root, "src", "setup.py")

                    app.indexing_jobs[job_id]["message"] = f"Running setup.py for {indexing_mode} indexing..."
                    app.indexing_jobs[job_id]["progress"] = 10

                    # Use venv Python instead of system Python to access installed packages
                    venv_python = os.path.join(chatbot_root, "venv", "Scripts", "python.exe")
                    python_cmd = venv_python if os.path.exists(venv_python) else sys.executable

                    if index_type == "sitemap" and sitemap:
                        # Sitemap: Index only sitemap URLs, skip PDFs
                        cmd = [python_cmd, setup_path, "-n", client_id, "-w", "-s", sitemap]
                        log_info(f"[Index] Sitemap mode for {client_id} sitemap={sitemap}")
                    elif index_type == "website" and urls:
                        # Custom URLs: Index only specified URLs, skip PDFs
                        cmd = [python_cmd, setup_path, "-n", client_id, "-w", "-u", urls]
                        log_info(f"[Index] URL mode for {client_id} urls={urls}")
                    elif index_type == "website":
                        # Full Website: Index website recursively + PDFs (no -w flag!)
                        cmd = [python_cmd, setup_path, "-n", client_id]
                        log_info(f"[Index] Full website mode for {client_id} (includes PDFs)")
                    elif index_type == "pdf":
                        # PDF Only: Index only PDFs, skip website
                        cmd = [python_cmd, setup_path, "-n", client_id, "-p"]
                        log_info(f"[Index] PDF mode for {client_id} (PDFs only, no website)")
                    else:
                        # All/Default: Index both website + PDFs
                        cmd = [python_cmd, setup_path, "-n", client_id]
                        log_info(f"[Index] Full (default) mode for {client_id} (website + PDFs)")

                    log_info(f"[Index] Running command: {' '.join(cmd)}")
                    result = subprocess.run(
                        cmd,
                        cwd=chatbot_root,
                        capture_output=True,
                        text=True,
                    )

                    log_info(f"[Index] Command return code: {result.returncode}")
                    if result.stdout:
                        log_info(f"[Index] Command stdout: {result.stdout[:500]}...")
                    if result.stderr:
                        log_error(f"[Index] Command stderr: {result.stderr[:500]}...")

                    if result.returncode == 0:
                        app.indexing_jobs[job_id].update(
                            {
                                "status": "completed",
                                "progress": 100,
                                "message": f"{indexing_mode} indexing completed successfully",
                            }
                        )
                        log_info(f"{indexing_mode} indexing completed for {client_id}: {job_id}")
                        _save_indexing_history(client_id, indexing_mode_detail, urls, sitemap, "completed", client_configs)
                    else:
                        raise RuntimeError(result.stderr or "Unknown error")
                except Exception as exc_inner:
                    app.indexing_jobs[job_id].update(
                        {"status": "failed", "message": f"Indexing failed: {exc_inner}", "progress": 100}
                    )
                    log_error(f"Indexing error for {client_id}: {exc_inner}")
                    _save_indexing_history(client_id, indexing_mode_detail, urls, sitemap, "failed", client_configs)

            threading.Thread(target=run_indexing, daemon=True).start()
            log_info(f"{indexing_mode} indexing started for {client_id}: {job_id}")

            # Save initial history entry
            _save_indexing_history(client_id, indexing_mode_detail, urls, sitemap, "started", client_configs)

            return jsonify({"message": "Indexing job started", "job_id": job_id, "status": "processing"}), 202
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except Exception as exc:
            log_error(f"Error starting indexing: {exc}")
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/indexing/status/<job_id>", methods=["GET"])
    def indexing_status(job_id):
        try:
            jobs = getattr(app, "indexing_jobs", {})
            log_info(f"[Status] Checking job {job_id}, total jobs: {len(jobs)}, job exists: {job_id in jobs}")
            if job_id not in jobs:
                log_error(f"[Status] Job {job_id} not found in jobs: {list(jobs.keys())}")
                return jsonify({"error": "Job not found"}), 404
            return jsonify({"job_id": job_id, **jobs[job_id]}), 200
        except Exception as exc:
            log_error(f"Error checking indexing status: {exc}")
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/indexing/history", methods=["GET"])
    def get_indexing_history():
        try:
            client_id = request.args.get("client_id")
            config = _validate_client(client_id)
            root_dir = _client_root(config, client_id)

            # Read history file
            history_file = os.path.join(root_dir, client_id, "indexing_history.json")
            if os.path.exists(history_file):
                with open(history_file, 'r', encoding='utf-8') as f:
                    history = json.load(f)
            else:
                history = []

            # Filter to only show completed indexing operations
            completed_history = [entry for entry in history if entry.get("status") == "completed"]

            # Return most recent 20 completed entries
            return jsonify({"history": completed_history[-20:]}), 200

        except Exception as exc:
            log_error(f"Error retrieving indexing history: {exc}")
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/indexing/clear-vectorstore", methods=["POST"])
    def clear_vectorstore():
        """
        Clear the vectorstore to allow fresh re-indexing.
        Used when user wants to "Replace" existing indexed content.
        """
        import shutil
        try:
            data = request.json or {}
            client_id = data.get("client_id")
            
            config = _validate_client(client_id)
            root_dir = _client_root(config, client_id)
            
            # Get vectorstore path from config
            vector_store_file = config.get("VECTOR_STORE_FILE", "vectorstore.db")
            vectorstore_path = os.path.join(root_dir, client_id, vector_store_file)
            
            # Clear the vectorstore directory if it exists
            if os.path.exists(vectorstore_path):
                # Remove all files in the vectorstore directory
                for filename in os.listdir(vectorstore_path):
                    file_path = os.path.join(vectorstore_path, filename)
                    try:
                        if os.path.isfile(file_path):
                            os.remove(file_path)
                        elif os.path.isdir(file_path):
                            shutil.rmtree(file_path)
                    except Exception as e:
                        log_error(f"Error removing {file_path}: {e}")
                
                log_info(f"Cleared vectorstore for {client_id}")
                return jsonify({
                    "success": True,
                    "message": f"Vectorstore cleared for {client_id}",
                    "path": vectorstore_path
                }), 200
            else:
                return jsonify({
                    "success": True,
                    "message": "Vectorstore directory does not exist (nothing to clear)",
                    "path": vectorstore_path
                }), 200
                
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except Exception as exc:
            log_error(f"Error clearing vectorstore: {exc}")
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/indexing/check-duplicate", methods=["POST"])
    def check_duplicate_indexing():
        """
        Check if a sitemap URL or custom URLs have already been indexed.
        Returns duplicate info if found in indexing history.
        """
        try:
            data = request.json or {}
            client_id = data.get("client_id")
            index_type = (data.get("index_type") or "").strip()  # 'sitemap' or 'website'
            sitemap = (data.get("sitemap") or "").strip()
            urls = (data.get("urls") or "").strip()
            
            config = _validate_client(client_id)
            root_dir = _client_root(config, client_id)
            
            # Read history file
            history_file = os.path.join(root_dir, client_id, "indexing_history.json")
            if not os.path.exists(history_file):
                return jsonify({
                    "is_duplicate": False,
                    "duplicates": [],
                    "checked_value": sitemap or urls
                }), 200
            
            with open(history_file, 'r', encoding='utf-8') as f:
                history = json.load(f)
            
            # Filter to only completed entries
            completed_history = [entry for entry in history if entry.get("status") == "completed"]
            
            duplicates = []
            
            if index_type == "sitemap" and sitemap:
                # Check for matching sitemap URL
                sitemap_lower = sitemap.lower().strip()
                for entry in completed_history:
                    entry_sitemap = (entry.get("sitemap") or "").lower().strip()
                    if entry_sitemap == sitemap_lower:
                        duplicates.append({
                            "type": "sitemap",
                            "value": entry.get("sitemap"),
                            "timestamp": entry.get("timestamp"),
                            "indexing_mode": entry.get("indexing_mode")
                        })
            
            elif urls:
                # Check for matching URLs
                # Split the incoming URLs and normalize
                url_list = [u.strip().lower() for u in urls.split(",") if u.strip()]
                
                for entry in completed_history:
                    entry_urls = entry.get("urls", "")
                    if entry_urls:
                        entry_url_list = [u.strip().lower() for u in entry_urls.split(",") if u.strip()]
                        # Check for any overlap
                        matching_urls = set(url_list) & set(entry_url_list)
                        if matching_urls:
                            duplicates.append({
                                "type": "urls",
                                "value": entry.get("urls"),
                                "matching_urls": list(matching_urls),
                                "timestamp": entry.get("timestamp"),
                                "indexing_mode": entry.get("indexing_mode")
                            })
            
            log_info(f"Indexing duplicate check for {index_type}: found {len(duplicates)} duplicate(s)")
            return jsonify({
                "is_duplicate": len(duplicates) > 0,
                "duplicates": duplicates,
                "index_type": index_type,
                "checked_value": sitemap or urls
            }), 200
            
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except Exception as exc:
            log_error(f"Error checking for indexing duplicates: {exc}")
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/keys", methods=["GET"])
    def list_api_keys():
        try:
            client_id = request.args.get("client_id")
            config = _validate_client(client_id)
            path = _ensure_secrets_path(_client_root(config, client_id), client_id)
            store = _load_key_store(path)

            keys = []
            for provider_key, entry in store.items():
                keys.append(
                    {
                        "client_id": client_id,
                        "provider": entry.get("provider") or provider_key,
                        "key_name": entry.get("key_name") or (entry.get("provider") or provider_key).title(),
                        "masked_key": _mask_key(entry.get("key_value")),
                        "updated_at": entry.get("updated_at"),
                        "env_var": entry.get("env_var") or API_PROVIDER_ENV_MAP.get(provider_key),
                    }
                )

            return jsonify({"keys": keys, "total": len(keys)}), 200
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except Exception as exc:
            log_error(f"Error listing API keys: {exc}")
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/keys", methods=["POST"])
    def upsert_api_key():
        try:
            payload = request.get_json() or {}
            client_id = payload.get("client_id")
            provider = (payload.get("provider") or "").strip()
            key_value = (payload.get("key_value") or "").strip().strip('"').strip("'")
            key_name = (payload.get("key_name") or provider.title()).strip()
            env_var = (payload.get("env_var") or "").strip()

            if not client_id or client_id not in client_configs:
                return jsonify({"error": "Invalid client_id"}), 400
            if not provider or not key_value:
                return jsonify({"error": "provider and key_value are required"}), 400

            provider_key = provider.lower()
            config = client_configs[client_id]
            path = _ensure_secrets_path(_client_root(config, client_id), client_id)
            store = _load_key_store(path)

            entry = {
                "provider": provider,
                "key_name": key_name or provider.title(),
                "key_value": key_value,
                "env_var": env_var or API_PROVIDER_ENV_MAP.get(provider_key),
                "updated_at": datetime.utcnow().isoformat(),
            }
            store[provider_key] = entry
            _save_key_store(path, store)
            _apply_env(provider_key, entry, logger)

            log_info(f"Stored API key for {client_id}:{provider}")
            return (
                jsonify(
                    {
                        "client_id": client_id,
                        "provider": entry["provider"],
                        "key_name": entry["key_name"],
                        "masked_key": _mask_key(key_value),
                        "updated_at": entry["updated_at"],
                        "env_var": entry["env_var"],
                    }
                ),
                201,
            )
        except Exception as exc:
            log_error(f"Error saving API key: {exc}")
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/keys/<provider>", methods=["DELETE"])
    def delete_api_key(provider):
        try:
            client_id = request.args.get("client_id")
            config = _validate_client(client_id)
            provider_key = provider.lower()
            path = _ensure_secrets_path(_client_root(config, client_id), client_id)
            store = _load_key_store(path)

            entry = store.pop(provider_key, None)
            if not entry:
                return jsonify({"error": "Key not found"}), 404

            _save_key_store(path, store)
            log_info(f"Removed API key for {client_id}:{provider_key}")
            return jsonify({"message": "API key deleted", "provider": entry.get("provider", provider_key)}), 200
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except Exception as exc:
            log_error(f"Error deleting API key: {exc}")
            return jsonify({"error": str(exc)}), 500

    # ========== AUDIT LOG ENDPOINTS (Database-based using Log.db) ==========
    @app.route("/api/audit/logs", methods=["GET"])
    def list_audit_logs():
        """List conversation audit logs for a client from SQLite database"""
        import sqlite3
        try:
            client_id = request.args.get("client_id")
            config = _validate_client(client_id)
            limit = int(request.args.get("limit", 50))
            offset = int(request.args.get("offset", 0))
            search = request.args.get("search", "")

            # Get database path - try multiple locations
            log_db_path = config.get("LOG_DB_PATH")
            if not log_db_path:
                root_dir = _client_root(config, client_id)
                # Try relative to current directory first (most common)
                log_db_path = os.path.join("application_db", "log_db")
                if not os.path.exists(os.path.join(log_db_path, "Log.db")):
                    # Try under Data directory
                    log_db_path = os.path.join(root_dir, "application_db", "log_db")

            log_db_file = os.path.join(log_db_path, "Log.db")

            if not os.path.exists(log_db_file):
                log_error(f"Database not found at: {log_db_file}")
                return jsonify({"logs": [], "total": 0, "limit": limit, "offset": offset}), 200

            # Query database
            conn = sqlite3.connect(log_db_file)
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()

            # Build query with search filter
            where_clause = "WHERE client_id = ?"
            params = [client_id]

            if search:
                where_clause += " AND (session_id LIKE ? OR conversation LIKE ?)"
                search_param = f"%{search}%"
                params.extend([search_param, search_param])

            # Get total count
            cursor.execute(f"SELECT COUNT(*) as total FROM client_sessions {where_clause}", params)
            total = cursor.fetchone()["total"]

            # Get paginated results
            query = f"""
            SELECT id, client_id, session_id, conversation, time, date, summary
            FROM client_sessions
            {where_clause}
            ORDER BY id DESC
            LIMIT ? OFFSET ?
            """
            params.extend([limit, offset])
            cursor.execute(query, params)
            rows = cursor.fetchall()

            all_logs = []
            for row in rows:
                conversation = json.loads(row["conversation"]) if row["conversation"] else {}

                # Extract last question and response
                last_question = None
                last_response = None
                turns = len(conversation)

                # Filter out conversations with less than 3 turns (user messages)
                # This matches the filter in conv_summarizer.py which requires > 5 total messages
                # 3 turns ≈ 3 user + 3 AI = 6 total messages (meets the > 5 requirement)
                if turns < 3:
                    continue

                if conversation:
                    last_key = list(conversation.keys())[-1] if conversation.keys() else None
                    if last_key:
                        last_question = last_key
                        last_entry = conversation[last_key]
                        if isinstance(last_entry, dict):
                            last_response = last_entry.get("response")

                all_logs.append({
                    "id": row["id"],
                    "client_id": row["client_id"],
                    "session_id": row["session_id"],
                    "date": row["date"],
                    "time": row["time"],
                    "summary": bool(row["summary"]),
                    "turns": turns,
                    "last_question": last_question,
                    "last_response": last_response,
                    "updated_at": f"{row['date']} {row['time']}"
                })

            conn.close()

            return jsonify({
                "logs": all_logs,
                "total": total,
                "limit": limit,
                "offset": offset
            }), 200

        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except Exception as exc:
            log_error(f"Error fetching audit logs: {exc}")
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/audit/logs/<session_id>", methods=["GET"])
    def get_audit_log_detail(session_id):
        """Get detailed conversation for a specific session from SQLite database"""
        import sqlite3
        try:
            client_id = request.args.get("client_id")
            config = _validate_client(client_id)

            # Get database path - try multiple locations
            log_db_path = config.get("LOG_DB_PATH")
            if not log_db_path:
                root_dir = _client_root(config, client_id)
                # Try relative to current directory first (most common)
                log_db_path = os.path.join("application_db", "log_db")
                if not os.path.exists(os.path.join(log_db_path, "Log.db")):
                    # Try under Data directory
                    log_db_path = os.path.join(root_dir, "application_db", "log_db")

            log_db_file = os.path.join(log_db_path, "Log.db")

            if not os.path.exists(log_db_file):
                log_error(f"Database not found at: {log_db_file}")
                return jsonify({"error": "Database not found"}), 404

            # Query database for specific session
            conn = sqlite3.connect(log_db_file)
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()

            cursor.execute("""
            SELECT id, client_id, session_id, conversation, time, date, summary
            FROM client_sessions
            WHERE session_id = ? AND client_id = ?
            """, (session_id, client_id))

            row = cursor.fetchone()
            conn.close()

            if not row:
                return jsonify({"error": "Session not found"}), 404

            conversation = json.loads(row["conversation"]) if row["conversation"] else {}

            # Convert conversation to pairs format
            conversation_pairs = []
            for question, data in conversation.items():
                if isinstance(data, dict):
                    conversation_pairs.append({
                        "question": question,
                        "response": data.get("response"),
                        "metadata": {
                            "options": data.get("options", []),
                            "chatMessageOptions": data.get("chatMessageOptions", []),
                            "jobs": data.get("jobs", [])
                        }
                    })

            return jsonify({
                "id": row["id"],
                "client_id": row["client_id"],
                "session_id": row["session_id"],
                "date": row["date"],
                "time": row["time"],
                "summary": bool(row["summary"]),
                "turns": len(conversation),
                "conversation": conversation,
                "conversation_pairs": conversation_pairs,
                "updated_at": f"{row['date']} {row['time']}"
            }), 200

        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except Exception as exc:
            log_error(f"Error fetching audit log detail: {exc}")
            return jsonify({"error": str(exc)}), 500

    # ========== APPLICATION LOGS ENDPOINTS ==========
    
    def _get_log_path(config: Dict) -> str:
        """Get the application logs path from config or default."""
        # Check for APPLICATION_LOG_PATH in config, default to 'application_logs'
        log_path = config.get("APPLICATION_LOG_PATH", "application_logs")
        # If it's a relative path, make it relative to the chatbot's root
        if not os.path.isabs(log_path):
            chatbot_root = config.get("CHATBOT_ROOT", os.getcwd())
            log_path = os.path.join(chatbot_root, log_path)
        return log_path

    @app.route("/api/logs", methods=["GET"])
    def list_application_logs():
        """List all application log files"""
        try:
            client_id = request.args.get("client_id")
            config = _validate_client(client_id)
            
            log_path = _get_log_path(config)
            
            if not os.path.exists(log_path):
                return jsonify({"logs": [], "total": 0, "log_path": log_path}), 200
            
            log_files = []
            # Get all .log files in the directory
            for filename in os.listdir(log_path):
                if filename.endswith(".log") or filename.endswith(".txt"):
                    filepath = os.path.join(log_path, filename)
                    try:
                        stat = os.stat(filepath)
                        file_size = stat.st_size
                        modified_time = datetime.fromtimestamp(stat.st_mtime)
                        created_time = datetime.fromtimestamp(stat.st_ctime)
                        
                        # Calculate age in days
                        age_days = (datetime.now() - modified_time).days
                        
                        log_files.append({
                            "filename": filename,
                            "filepath": filepath,
                            "size": file_size,
                            "size_formatted": _format_size(file_size),
                            "modified_at": modified_time.isoformat(),
                            "created_at": created_time.isoformat(),
                            "age_days": age_days,
                            "is_old": age_days >= 7
                        })
                    except Exception as e:
                        log_error(f"Error reading log file {filename}: {e}")
                        continue
            
            # Sort by modified time descending (newest first)
            log_files.sort(key=lambda x: x["modified_at"], reverse=True)
            
            return jsonify({
                "logs": log_files,
                "total": len(log_files),
                "log_path": log_path,
                "old_logs_count": sum(1 for f in log_files if f["is_old"])
            }), 200
            
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except Exception as exc:
            log_error(f"Error listing application logs: {exc}")
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/logs/<filename>", methods=["GET"])
    def get_application_log(filename):
        """Get contents of a specific log file"""
        try:
            client_id = request.args.get("client_id")
            config = _validate_client(client_id)
            
            log_path = _get_log_path(config)
            filepath = os.path.join(log_path, filename)
            
            # Security check: ensure the file is within the log directory
            if not os.path.realpath(filepath).startswith(os.path.realpath(log_path)):
                return jsonify({"error": "Invalid file path"}), 400
            
            if not os.path.exists(filepath):
                return jsonify({"error": "Log file not found"}), 404
            
            # Get optional parameters for pagination
            offset = int(request.args.get("offset", 0))
            limit = int(request.args.get("limit", 1000))  # Default 1000 lines
            tail = request.args.get("tail", "false").lower() == "true"
            
            try:
                with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                    all_lines = f.readlines()
                
                total_lines = len(all_lines)
                
                if tail:
                    # Return last 'limit' lines
                    lines = all_lines[-limit:] if limit < total_lines else all_lines
                    offset = max(0, total_lines - limit)
                else:
                    # Return lines from offset
                    lines = all_lines[offset:offset + limit]
                
                # Get file metadata
                stat = os.stat(filepath)
                
                return jsonify({
                    "filename": filename,
                    "content": "".join(lines),
                    "lines": lines,
                    "total_lines": total_lines,
                    "offset": offset,
                    "limit": limit,
                    "size": stat.st_size,
                    "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat()
                }), 200
                
            except Exception as e:
                log_error(f"Error reading log file {filename}: {e}")
                return jsonify({"error": f"Failed to read log file: {str(e)}"}), 500
            
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except Exception as exc:
            log_error(f"Error getting application log: {exc}")
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/logs/<filename>", methods=["DELETE"])
    def delete_application_log(filename):
        """Delete a specific log file"""
        try:
            client_id = request.args.get("client_id")
            config = _validate_client(client_id)
            
            log_path = _get_log_path(config)
            filepath = os.path.join(log_path, filename)
            
            # Security check: ensure the file is within the log directory
            if not os.path.realpath(filepath).startswith(os.path.realpath(log_path)):
                return jsonify({"error": "Invalid file path"}), 400
            
            if not os.path.exists(filepath):
                return jsonify({"error": "Log file not found"}), 404
            
            # Check if this is the active log file (likely locked by logger)
            # Try to delete, handle locked file gracefully
            try:
                os.remove(filepath)
                log_info(f"Deleted application log file: {filename} for client {client_id}")
                return jsonify({
                    "message": "Log file deleted successfully",
                    "filename": filename
                }), 200
            except PermissionError:
                # File is locked (likely active log file) - try to truncate instead
                try:
                    with open(filepath, 'w') as f:
                        f.truncate(0)
                    log_info(f"Cleared (truncated) active log file: {filename} for client {client_id}")
                    return jsonify({
                        "message": "Log file cleared (file was in use, contents truncated)",
                        "filename": filename,
                        "note": "Active log file was cleared instead of deleted"
                    }), 200
                except Exception as truncate_error:
                    return jsonify({
                        "error": f"Cannot delete or clear log file - it's currently in use by the application. Try again after restarting the chatbot.",
                        "filename": filename
                    }), 423  # 423 Locked
            
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except Exception as exc:
            log_error(f"Error deleting application log: {exc}")
            return jsonify({"error": str(exc)}), 500

    @app.route("/api/logs/cleanup", methods=["POST"])
    def cleanup_old_logs():
        """Delete log files older than specified days (default 7 days)"""
        try:
            data = request.get_json() or {}
            client_id = data.get("client_id") or request.args.get("client_id")
            config = _validate_client(client_id)
            
            days = int(data.get("days", 7))
            log_path = _get_log_path(config)
            
            if not os.path.exists(log_path):
                return jsonify({
                    "message": "No logs directory found",
                    "deleted_count": 0,
                    "deleted_files": []
                }), 200
            
            cutoff_date = datetime.now() - timedelta(days=days)
            deleted_files = []
            
            skipped_files = []
            for filename in os.listdir(log_path):
                if not (filename.endswith(".log") or filename.endswith(".txt")):
                    continue
                    
                filepath = os.path.join(log_path, filename)
                try:
                    modified_time = datetime.fromtimestamp(os.path.getmtime(filepath))
                    if modified_time < cutoff_date:
                        try:
                            os.remove(filepath)
                            deleted_files.append(filename)
                            log_info(f"Cleaned up old log file: {filename}")
                        except PermissionError:
                            # File is locked, try to truncate
                            try:
                                with open(filepath, 'w') as f:
                                    f.truncate(0)
                                deleted_files.append(f"{filename} (cleared)")
                                log_info(f"Cleared old locked log file: {filename}")
                            except Exception:
                                skipped_files.append(filename)
                                log_info(f"Skipped locked log file: {filename}")
                except Exception as e:
                    log_error(f"Error processing old log file {filename}: {e}")
                    continue
            
            log_info(f"Log cleanup completed for {client_id}: {len(deleted_files)} files deleted, {len(skipped_files)} skipped")
            
            message = f"Cleanup completed. Deleted/cleared {len(deleted_files)} log files older than {days} days."
            if skipped_files:
                message += f" {len(skipped_files)} file(s) skipped (in use)."
            
            return jsonify({
                "message": message,
                "deleted_count": len(deleted_files),
                "deleted_files": deleted_files,
                "skipped_files": skipped_files,
                "days_threshold": days
            }), 200
            
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except Exception as exc:
            log_error(f"Error cleaning up logs: {exc}")
            return jsonify({"error": str(exc)}), 500

    def _format_size(size_bytes: int) -> str:
        """Format file size in human readable format"""
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        elif size_bytes < 1024 * 1024 * 1024:
            return f"{size_bytes / (1024 * 1024):.1f} MB"
        else:
            return f"{size_bytes / (1024 * 1024 * 1024):.1f} GB"

    log_info("Admin API endpoints registered successfully (including application logs)")